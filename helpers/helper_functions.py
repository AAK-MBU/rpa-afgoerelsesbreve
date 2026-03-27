"""
Utility functions used by the skabelonmotor Excel parser.
"""

import base64
import copy
import logging
import os
import re
import urllib.parse

from datetime import datetime
from io import BytesIO

from docx import Document

from openpyxl import load_workbook
from openpyxl.cell.rich_text import CellRichText

import pandas as pd

from sqlalchemy import create_engine, text

logger = logging.getLogger(__name__)

# Regex used to detect block headers such as:
# "Blok 1", "Blok 3.1", "Blok 7.2a"
BLOCK_HEADER_PATTERN = re.compile(r"^Blok\s+([0-9]+(?:\.\s*[0-9]+)?[a-zA-Z]?)")


def resolve_blocks(blocks: list[dict], block_metadata: dict, item_data: dict):
    """
    Some blocks in the template block data need a custom key or function or similar to be able to be properly handled by the skabelonmotor.
    This helper function is responsible for looping through the retrieved template block data, and handling custom keys and similar requirements.
    """

    blocks = copy.deepcopy(blocks)

    # We loop through each block and check if the block_id is in any of the specified custom handlers
    for block in blocks:
        block_id = block.get("block_id")

        if not block_id:
            continue

        # -------------------------
        # DEFAULT
        # -------------------------
        block["condition"] = "equals"

        # -------------------------
        # COPY
        # If the block_id is specified in the copy section, we copy the mapping, entries, and condition from the specified copy block
        # This should always be from a block that is handled earlier than the block specified in the copy section of the block_metadata
        # -------------------------
        if block_id in block_metadata.get("copy", {}):
            source_id = block_metadata["copy"][block_id]

            source_block = next(
                (b for b in blocks if b["block_id"] == source_id),
                None
            )

            if source_block:
                block["mapping"] = source_block.get("mapping")
                block["entries"] = copy.deepcopy(source_block.get("entries", {}))
                block["condition"] = source_block.get("condition", "equals")

            continue

        # -------------------------
        # CUSTOM (function)
        # -------------------------
        if block_id in block_metadata.get("custom", {}):
            func = block_metadata["custom"][block_id]

            updated_block = func(item_data, block)

            block.update(updated_block)
            block["condition"] = "custom"

            continue

        # -------------------------
        # CUSTOM KEY (precomputed value)
        # -------------------------
        if block_id in block_metadata.get("custom_key", {}):
            value = block_metadata["custom_key"][block_id]

            if value:
                block["mapping"] = value
                block["condition"] = "custom"

            continue

        # -------------------------
        # HAS VALUE
        # -------------------------
        if block_id in block_metadata.get("has_value", []):
            block["condition"] = "has_value"

            continue

        # -------------------------
        # ALL
        # -------------------------
        if block_id in block_metadata.get("all", []):
            block["condition"] = "all"

            continue

    return blocks


def get_db_connection_string():
    """
    Database helper to retrieve the database connection string
    """

    return os.getenv("DBCONNECTIONSTRINGDEV")


def read_sql(query: str = "", params: dict = None, conn_string: str = "") -> pd.DataFrame:
    """
    Run a SELECT sql statement
    """

    if params is None:
        params = {}

    encoded_conn_str = urllib.parse.quote_plus(conn_string)

    engine = create_engine(f"mssql+pyodbc:///?odbc_connect={encoded_conn_str}")

    try:
        with engine.begin() as conn:
            df = pd.read_sql(text(query), conn, params=params)

        return df

    except Exception as e:
        logger.info(f"SQL error: {e}")

        raise


def replace_template_placeholders(template_bytes: str, data: dict) -> bytes:
    """
    Replaces all {{placeholders}} in a DOCX template with values from `data`.

    The function parses the document structure (paragraphs, tables, headers, footers)
    and safely replaces placeholders without breaking formatting, images, or Word fields.
    It handles cases where placeholders are split across multiple runs (a common Word behavior).
    Returns the updated document as a base64-encoded string.
    """

    doc = Document(BytesIO(template_bytes))

    # Normalize keys so template placeholders and data keys match consistently
    normalized_data = {
        normalize_key(k): str(v)
        for k, v in data.items()
        if v is not None
    }

    def replace_in_paragraph(paragraph):
        """
        Replaces placeholders inside a single paragraph.

        Word may split a placeholder like {{key}} across multiple runs,
        so we merge consecutive runs when detecting '{{' until '}}' is found.
        This ensures correct replacement while preserving formatting and embedded elements.
        """

        runs = paragraph.runs
        i = 0

        # Iterate through runs manually so we can merge forward when needed
        while i < len(runs):

            # Detect start of a placeholder
            if "{{" in runs[i].text:

                full_text = runs[i].text
                j = i

                # Merge subsequent runs until we find the closing '}}'
                while "}}" not in full_text and j + 1 < len(runs):
                    j += 1
                    full_text += runs[j].text

                # Find all placeholders inside the merged text
                matches = re.findall(r"\{\{(.*?)\}\}", full_text)

                for match in matches:
                    normalized_placeholder = normalize_key(match)

                    # Replace only if we have matching data
                    if normalized_placeholder in normalized_data:
                        value = normalized_data[normalized_placeholder]
                        full_text = full_text.replace(f"{{{{{match}}}}}", value)

                # Write updated text back to the FIRST run
                runs[i].text = full_text

                # Clear the remaining merged runs to avoid duplicate content
                for k in range(i + 1, j + 1):
                    runs[k].text = ""

                # Skip ahead to avoid reprocessing merged runs
                i = j

            i += 1

    def replace_in_table(table):
        """
        Recursively replaces placeholders inside tables.

        Word documents often store layout (especially headers) inside tables,
        including nested tables. This function ensures all cells and nested
        structures are processed so no placeholders are missed.
        """

        for row in table.rows:
            for cell in row.cells:

                # Process normal paragraphs inside the cell
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

                # Recursively process nested tables (Word supports this)
                for nested_table in cell.tables:
                    replace_in_table(nested_table)

    # -------------------------
    # Body paragraphs
    # -------------------------
    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    # -------------------------
    # Body tables
    # -------------------------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

    # -------------------------
    # Headers (including first page header)
    # -------------------------
    for section in doc.sections:

        # Standard header
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph)

        for table in section.header.tables:
            replace_in_table(table)

        # First-page header (used when "different first page" is enabled)
        for paragraph in section.first_page_header.paragraphs:
            replace_in_paragraph(paragraph)

        for table in section.first_page_header.tables:
            replace_in_table(table)

    # -------------------------
    # Footers
    # -------------------------
    for section in doc.sections:
        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph)

    # Save modified document to memory buffer
    buffer = BytesIO()
    doc.save(buffer)

    # Return as base64 (useful for APIs / transport)
    template_b64 = base64.b64encode(buffer.getvalue()).decode()

    return template_b64


def parse_date(value: str | None):
    """
    Convert a DD-MM-YYYY string into a datetime object.

    Used primarily for sorting transport rows by start/end date.

    Args:
        value (str | None): Date string in DD-MM-YYYY format.

    Returns:
        datetime: Parsed date or datetime.max if value is missing.
    """

    if not value:
        return datetime.max

    return datetime.strptime(value, "%d-%m-%Y")


def extract_cell_formatting(cell):
    """
    Convert Excel rich text content into HTML-like formatted text.

    Handles formatting such as bold, italic, underline, strike-through
    and color while preserving the original text structure.

    Args:
        cell: openpyxl cell object.

    Returns:
        str: HTML-like formatted text.
    """

    if cell is None or cell.value is None:
        return ""

    value = cell.value

    # ----------------------------------------
    # Rich formatted text (Excel rich text)
    # ----------------------------------------
    if isinstance(value, CellRichText):

        parts = []

        for block in value:
            text = block.text or ""
            font = block.font

            if not text:
                continue

            # Remove zero-width characters sometimes inserted by Excel
            text = text.replace("\u200b", "")

            # Replace Excel tab indentation
            text = text.replace("\t", " ")

            prefix = ""
            suffix = ""

            if font:
                # Bold
                if font.b:
                    prefix += "<strong>"
                    suffix = "</strong>" + suffix

                # Italic
                if font.i:
                    prefix += "<em>"
                    suffix = "</em>" + suffix

                # Underline
                if font.u in ["single", "double", "singleAccounting", "doubleAccounting", True]:
                    prefix += "<u>"
                    suffix = "</u>" + suffix

                # Strikethrough
                if font.strike:
                    prefix += "<strike>"
                    suffix = "</strike>" + suffix

                # Text color
                if font.color and font.color.rgb:
                    rgb = font.color.rgb[-6:]

                    # Skip default black text
                    if rgb != "000000":
                        prefix += f'<span style="color:#{rgb}">'
                        suffix = "</span>" + suffix

            parts.append(f"{prefix}{text}{suffix}")

        return "".join(parts)

    # ----------------------------------------
    # Plain text cell (no formatting)
    # ----------------------------------------
    return str(value)


def parse_workbook(binary_excel: bytes) -> list[dict]:
    """
    Pure Excel parser.

    Extracts blocks and their entries from the workbook without applying
    any business logic, metadata, or custom functions.

    Args:
        binary_excel (bytes): Excel workbook content.

    Returns:
        list[dict]: Raw extracted block structures.
    """

    wb = load_workbook(BytesIO(binary_excel), rich_text=True)

    parsed_blocks = []
    current_block = None

    # ----------------------------------------
    # Parse workbook sheets
    # ----------------------------------------
    for sheet_name in wb.sheetnames:

        if not sheet_name.startswith("Blok"):
            continue

        ws = wb[sheet_name]
        rows = list(ws.iter_rows())

        for i, row in enumerate(rows):

            col_a_cell = row[0] if len(row) > 0 else None
            col_b_cell = row[1] if len(row) > 1 else None

            col_a = col_a_cell.value if col_a_cell else None
            col_b = extract_cell_formatting(col_b_cell) if col_b_cell else None

            # ----------------------------------------
            # Detect block header
            # ----------------------------------------
            if isinstance(col_a, str):

                match = BLOCK_HEADER_PATTERN.match(col_a)

                if match:

                    block_id = match.group(1).replace(" ", "").strip()

                    # Mapping key from column C in next row
                    next_row = rows[i + 1] if i + 1 < len(rows) else None
                    next_col_c = None

                    if next_row and len(next_row) > 2:
                        next_col_c = next_row[2].value

                    current_block = {
                        "block_id": block_id,
                        "title": col_a,
                        "mapping": str(next_col_c).strip() if next_col_c else None,
                        "entries": {}
                    }

                    parsed_blocks.append(current_block)

                    continue

            if not current_block:
                continue

            # ----------------------------------------
            # Parse entries
            # ----------------------------------------
            if col_a and col_b:

                entry_text = col_b.strip()

                # Skip "Ingen tekst"
                if normalize_key(entry_text) == "ingentekst":
                    continue

                key = str(col_a)

                current_block["entries"][key] = entry_text

    return parsed_blocks


def parse_workbook_old(citizen_data: dict, binary_excel: bytes, block_metadata: dict) -> list[dict]:
    """
    Parse the Excel template into structured block definitions used by the skabelonmotor.

    The workbook is scanned for rows starting with "Blok X", which define logical
    text blocks in the template. Each block is converted into a dictionary containing
    a block id, mapping key, condition type and its possible text entries.

    Parsing happens in three main stages:
    1. Build lookup tables from `block_metadata` so each block knows its behavior
       (e.g. equals, has_value, custom, copy).
    2. Iterate through sheets and rows to detect block headers and collect their
       entries from the Excel template.
    3. Post-process blocks by executing custom handlers and resolving "copy"
       blocks that inherit content from other blocks.

    Args:
        citizen_data (dict): Citizen data used by custom block handlers.
        binary_excel (bytes): Excel workbook content.
        block_metadata (dict): Configuration describing block conditions and handlers.

    Returns:
        list[dict]: Parsed block structures ready for the template engine.
    """

    wb = load_workbook(BytesIO(binary_excel), rich_text=True)

    parsed_blocks = []

    condition_lookup = {}
    custom_function_lookup = {}
    custom_key_lookup = {}
    copy_lookup = {}

    # ----------------------------------------
    # Build condition lookup tables
    # ----------------------------------------
    # Converts block_metadata into quick lookup dictionaries so each block can easily determine its behavior.
    for condition, blocks in block_metadata.items():

        if condition == "custom":

            for block_id, func in blocks.items():

                condition_lookup[block_id] = "custom"
                custom_function_lookup[block_id] = func

        elif condition == "custom_key":

            for block_id, value in blocks.items():

                condition_lookup[block_id] = "custom_key"
                custom_key_lookup[block_id] = value

        elif condition == "copy":

            for block_id, source_block in blocks.items():

                condition_lookup[block_id] = "copy"
                copy_lookup[block_id] = source_block

        else:

            for block_id in blocks:
                condition_lookup[block_id] = condition

    current_block = {}

    # ----------------------------------------
    # Parse workbook sheets
    # ----------------------------------------
    for sheet_name in wb.sheetnames:
        # Ignore sheets not starting with "Blok"
        if not sheet_name.startswith("Blok"):
            continue

        ws = wb[sheet_name]
        rows = list(ws.iter_rows())

        for i, row in enumerate(rows):

            col_a_cell = row[0] if len(row) > 0 else None
            col_b_cell = row[1] if len(row) > 1 else None

            col_a = col_a_cell.value if col_a_cell else None
            col_b = extract_cell_formatting(col_b_cell) if col_b_cell else None

            # ----------------------------------------
            # Detect block header
            # ----------------------------------------
            if isinstance(col_a, str):
                match = BLOCK_HEADER_PATTERN.match(col_a)

                if match:
                    # Finish previous block if it required custom processing
                    if current_block and current_block["condition"] == "custom":

                        func = custom_function_lookup.get(current_block["block_id"])

                        if func:
                            func(
                                citizen_data,
                                current_block
                            )

                    block_id = match.group(1).replace(" ", "").strip()

                    # Mapping key is defined in column C of the next row
                    next_row = rows[i + 1] if i + 1 < len(rows) else None
                    next_col_c = None

                    if next_row and len(next_row) > 2:
                        next_col_c = next_row[2].value

                    condition = condition_lookup.get(block_id, "equals")

                    mapping = str(next_col_c).strip() if next_col_c else None

                    # custom_key behaves like a custom block but with a predefined mapping value
                    if condition == "custom_key":
                        mapping = custom_key_lookup.get(block_id)
                        condition = "custom"

                    current_block = {
                        "block_id": block_id,
                        "title": col_a,
                        "mapping": mapping,
                        "condition": condition,
                        "entries": {}
                    }

                    parsed_blocks.append(current_block)

                    continue

            if not current_block:
                continue

            # ----------------------------------------
            # Parse block entries
            # ----------------------------------------
            if col_a and col_b:
                entry_text = col_b.strip()

                # Skip placeholder rows like "Ingen tekst"
                if normalize_key(entry_text) == "ingentekst":
                    continue

                key = str(col_a)

                current_block["entries"][key] = entry_text

        # ----------------------------------------
        # Finalize last block in sheet
        # ----------------------------------------
        if current_block and current_block["condition"] == "custom":
            func = custom_function_lookup.get(current_block["block_id"])

            if func:
                func(
                    citizen_data,
                    current_block
                )

    # ----------------------------------------
    # Apply "copy" block behavior
    # ----------------------------------------
    block_map = {b["block_id"]: b for b in parsed_blocks}

    for block in parsed_blocks:

        if block["condition"] == "copy":

            source_id = copy_lookup.get(block["block_id"])
            source_block = block_map.get(source_id)

            if source_block:

                block["mapping"] = source_block["mapping"]
                block["entries"] = copy.deepcopy(source_block["entries"])
                block["condition"] = source_block["condition"]

    return parsed_blocks


def normalize_key(value: str) -> str:
    """
    Normalize text for reliable key comparison.

    Removes whitespace, punctuation and Danish characters
    so template keys can be compared consistently.

    Args:
        value (str): Input string.

    Returns:
        str: Normalized key.
    """

    return (
        value.strip()
        .lower()
        .replace(" ", "")
        .replace(".", "")
        .replace("ø", "oe")
        .replace("å", "aa")
        .replace("æ", "ae")
        .replace("?", "")
        .replace("-", "")
        .replace("_", "")
    )
