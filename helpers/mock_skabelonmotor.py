"""
╔══════════════════════════════════════════════════════════════════════╗
║ 🔥 TEMPORARY MOCK - REMOVE WHEN api-skabelonmotor IS LIVE 🔥          ║
╠══════════════════════════════════════════════════════════════════════╣
║ This module is a self-contained, in-process copy of the               ║
║ `/letter_creation/create_letter` endpoint from the api-skabelonmotor  ║
║ service (app/api/letter_creation.py + app/utils/helper_functions.py). ║
║                                                                        ║
║ It exists ONLY so the RPA can generate letters without the API being  ║
║ dockerised and online. Once the API is deployed, delete this file and ║
║ restore the HTTP call in processes/process_item.py (see the marker    ║
║ comment there).                                                        ║
╚══════════════════════════════════════════════════════════════════════╝
"""

import base64
import io
import re
import tempfile

from io import BytesIO

from bs4 import BeautifulSoup

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import RGBColor, Pt

from docx2pdf import convert


def create_letter(
    data: dict,
    block_data: list,
    file_type: str,
    file_name: str,
    custom_key_overrides: dict | None = None,
    template_b64: str | None = None,
) -> bytes:
    """
    In-process replacement for the api-skabelonmotor `create_letter` endpoint.

    Builds the letter text from `block_data`, replaces placeholders and returns
    the rendered document as raw bytes (docx or pdf), mirroring what the API
    previously returned as an HTTP response body.
    """

    blocks = block_data
    overrides = custom_key_overrides or {}
    file_type = file_type.lower()

    # Initialize an empty list to contain each formatted and updated text part
    letter_parts = []

    for block in blocks:
        mapping = block.get("mapping")
        condition = block.get("condition")
        entries = block.get("entries", {})

        normalized_entries = {
            normalize_key(k): v
            for k, v in entries.items()
        }

        # -----------------------------
        # CONDITION: all
        # If it's an all condition, we append all the entries for the block
        # -----------------------------
        if condition == "all":
            for text in entries.values():
                letter_parts.append(text)

        # -----------------------------
        # CONDITION: has_value
        # has_value condition simply looks up the mapping_key and sees if there is a value - if there is, we append the single entry for the block
        # -----------------------------
        elif condition == "has_value":
            if mapping and data.get(normalize_key(mapping)):
                text = next(iter(entries.values()), None)

                if text:
                    letter_parts.append(text)

        # -----------------------------
        # CONDITION: custom
        # Uses the custom mapping directly as the entry key.
        # -----------------------------
        elif condition == "custom":
            if mapping:
                normalized_mapping = normalize_key(mapping)

                text = normalized_entries.get(normalized_mapping)

                if text:
                    letter_parts.append(text)

        # -----------------------------
        # CONDITION: custom_contains
        # Used when the entry key should be found inside the custom mapping value.
        # -----------------------------
        elif condition == "custom_contains":
            if mapping:
                normalized_mapping = normalize_key(mapping)

                matching_entry_keys = sorted(
                    normalized_entries.keys(),
                    key=len,
                    reverse=True
                )

                for normalized_entry_key in matching_entry_keys:
                    if normalized_entry_key in normalized_mapping:
                        text = normalized_entries[normalized_entry_key]

                        if text:
                            letter_parts.append(text)

                        break

        # -----------------------------
        # CONDITION: equals
        # equals condition is the default - we look for a key that matches in the provided data. Afterwards we use the value from the matched key, to look through the entry keys, for a key that matches the found value
        # -----------------------------
        elif condition == "equals":
            normalized_mapping = normalize_key(mapping)

            # 1️⃣ check overrides first
            key = overrides.get(normalized_mapping)

            # 2️⃣ fallback to data
            if key is None:
                key = data.get(normalized_mapping)

            if key:
                # ---------------------------------
                # CASE: multiple keys
                # ---------------------------------
                if isinstance(key, list):
                    for item in key:
                        normalized_item = normalize_key(item)

                        text = normalized_entries.get(normalized_item)

                        if text:
                            letter_parts.append(text)

                # ---------------------------------
                # CASE: single key
                # ---------------------------------
                else:
                    normalized_item = normalize_key(key)

                    text = normalized_entries.get(normalized_item)

                    if text:
                        letter_parts.append(text)

    # ---------------------------------
    # Combine blocks and replace placeholders
    # ---------------------------------
    letter_text = "\n\n".join(letter_parts)
    letter_text = replace_placeholders(letter_text, data)

    text = normalize_html(text=letter_text)

    # Here we check if the request included a docx template
    # If it did, we simply insert the letter_text into that template - if not, we must create the docx from scratch
    if template_b64:
        docx_bytes = insert_letter_into_template(template_b64=template_b64, letter_text=text)

    else:
        docx_bytes = html_to_docx_bytes(text=letter_text)

    if file_type == "docx":
        return docx_bytes

    if file_type == "pdf":
        # Because we always, by default, create the letter as a Word docx, we must convert it to pdf if necessary
        return convert_docx_to_pdf(docx_bytes)

    raise ValueError(f"Unsupported file_type: {file_type}")


def add_hyperlink(paragraph, url, text):

    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")

    # Blue color
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    # Underline
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    new_run.append(rPr)

    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def insert_letter_into_template(template_b64: str, letter_text: str) -> bytes:
    """
    Insert rendered HTML letter text into a DOCX template.

    The function locates the {{LETTER_TEXT}} placeholder in the template,
    removes it, and inserts the formatted HTML content at the same location.
    """

    template_bytes = base64.b64decode(template_b64)

    doc = Document(BytesIO(template_bytes))

    # -------------------------------------------------
    # Recursive HTML → DOCX run processor
    # -------------------------------------------------
    def process_node(node, paragraph, formatting=None):

        if formatting is None:
            formatting = {}

        # ------------------------------
        # TEXT NODE
        # ------------------------------
        if node.name is None:

            content = str(node)

            if not content.strip():
                # Preserve intentional line breaks even when the newline sits
                # between inline tags (e.g. "...\n</em>\n<em>..."). Splitting the
                # letter on "\n\n" only catches *consecutive* newlines; a blank
                # line whose two newlines straddle a tag boundary arrives here as
                # a standalone "\n" text node and would otherwise be discarded,
                # collapsing the blank line to nothing.
                if "\n" in content:
                    run = paragraph.add_run()
                    for _ in range(content.count("\n")):
                        run.add_break()
                elif " " in content or "\xa0" in content:
                    paragraph.add_run(" ")

                return

            run = paragraph.add_run(content)

            if formatting.get("bold"):
                run.bold = True

            if formatting.get("italic"):
                run.italic = True

            if formatting.get("underline"):
                run.underline = True

            if formatting.get("strike"):
                run.font.strike = True

            rgb = formatting.get("color")

            if isinstance(rgb, str) and len(rgb) == 6:
                run.font.color.rgb = RGBColor.from_string(rgb)

        # ------------------------------
        # ELEMENT NODE
        # ------------------------------
        else:

            new_format = formatting.copy()

            if node.name == "a":

                url = node.get("href")

                # Create empty hyperlink first
                hyperlink_text = ""

                for child in node.children:
                    hyperlink_text += str(child)

                add_hyperlink(paragraph, url, hyperlink_text)

                return

            if node.name in ["strong", "b"]:
                new_format["bold"] = True

            if node.name in ["em", "i"]:
                new_format["italic"] = True

            if node.name == "u":
                new_format["underline"] = True

            if node.name == "strike":
                new_format["strike"] = True

            if node.name in ["span", "font"]:
                match = re.search(r"#([0-9A-Fa-f]{6})", str(node))
                if match:
                    new_format["color"] = match.group(1)

            # Recursively process child nodes so formatting cascades
            for child in node.children:
                process_node(child, paragraph, new_format)

    # -------------------------------------------------
    # Find placeholder and insert content
    # -------------------------------------------------
    for paragraph in doc.paragraphs:

        if "{{LETTER_TEXT}}" in paragraph.text.upper():

            parent = paragraph._element.getparent()
            index = parent.index(paragraph._element)

            # Remove placeholder paragraph
            parent.remove(paragraph._element)

            paragraphs = letter_text.split("\n\n")

            for offset, p in enumerate(paragraphs):

                new_paragraph = doc.add_paragraph()

                soup = BeautifulSoup(p, "html.parser")

                for child in soup.children:
                    process_node(child, new_paragraph)

                # Move paragraph to correct location
                parent.insert(index + offset, new_paragraph._element)

                new_paragraph.paragraph_format.space_after = Pt(12)

            break

    buffer = BytesIO()
    doc.save(buffer)

    return buffer.getvalue()


def normalize_html(text: str) -> str:
    """
    Normalize lightweight HTML formatting so it works with the rendering engines.
    """

    # ReportLab does not support <span style="color:..."> but does support <font>
    text = re.sub(
        r'<span style="color:#([0-9A-Fa-f]{6})">',
        r'<font color="#\1">',
        text
    )

    # Close converted color tags
    text = text.replace("</span>", "</font>")

    # ReportLab expects <b> and <i> rather than <strong> / <em>
    text = text.replace("<strong>", "<b>").replace("</strong>", "</b>")
    text = text.replace("<em>", "<i>").replace("</em>", "</i>")

    return text


def html_to_docx_bytes(text: str) -> bytes:
    """
    Render HTML-like formatted text into a DOCX document.
    """

    doc = Document()

    def process_node(node, paragraph, formatting=None):

        if formatting is None:
            formatting = {}

        # ----------------------------------------
        # TEXT NODE
        # ----------------------------------------
        if node.name is None:

            content = str(node)

            if not content.strip():
                # Preserve intentional line breaks even when the newline sits
                # between inline tags (e.g. "...\n</em>\n<em>..."). Splitting the
                # letter on "\n\n" only catches *consecutive* newlines; a blank
                # line whose two newlines straddle a tag boundary arrives here as
                # a standalone "\n" text node and would otherwise be discarded,
                # collapsing the blank line to nothing.
                if "\n" in content:
                    run = paragraph.add_run()
                    for _ in range(content.count("\n")):
                        run.add_break()
                elif " " in content or "\xa0" in content:
                    paragraph.add_run(" ")

                return

            run = paragraph.add_run(content)

            if formatting.get("bold"):
                run.bold = True

            if formatting.get("italic"):
                run.italic = True

            if formatting.get("underline"):
                run.underline = True

            if formatting.get("strike"):
                run.font.strike = True

            rgb = formatting.get("color")

            if isinstance(rgb, str) and len(rgb) == 6:
                run.font.color.rgb = RGBColor.from_string(rgb)

        else:

            # ----------------------------------------
            # ELEMENT NODE
            # ----------------------------------------
            new_format = formatting.copy()

            if node.name == "a":

                url = node.get("href")

                # Create empty hyperlink first
                hyperlink_text = ""

                for child in node.children:
                    hyperlink_text += str(child)

                add_hyperlink(paragraph, url, hyperlink_text)

                return

            if node.name in ["strong", "b"]:
                new_format["bold"] = True

            if node.name in ["em", "i"]:
                new_format["italic"] = True

            if node.name == "u":
                new_format["underline"] = True

            if node.name == "strike":
                new_format["strike"] = True

            if node.name in ["span", "font"]:
                match = re.search(r"#([0-9A-Fa-f]{6})", str(node))
                if match:
                    new_format["color"] = match.group(1)

            # Recursively process child nodes so formatting cascades
            for child in node.children:
                process_node(child, paragraph, new_format)

    # ----------------------------------------
    # Build DOCX paragraphs
    # ----------------------------------------
    # Paragraphs in the template engine are separated by double line breaks.
    paragraphs = text.split("\n\n")

    for p in paragraphs:

        paragraph = doc.add_paragraph()

        # Parse paragraph HTML so formatting can be processed node-by-node
        soup = BeautifulSoup(p, "html.parser")

        # Each top-level node inside the paragraph is processed
        for child in soup.children:
            process_node(child, paragraph)

    buffer = io.BytesIO()
    doc.save(buffer)

    return buffer.getvalue()


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """
    Helper function to convert a Word docx to pdf bytes
    """

    with tempfile.TemporaryDirectory() as tmpdir:

        docx_path = f"{tmpdir}/file.docx"
        pdf_path = f"{tmpdir}/file.pdf"

        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        convert(docx_path, pdf_path)

        with open(pdf_path, "rb") as f:
            return f.read()


def normalize_key(value: str) -> str:
    """
    Normalize strings for reliable key comparisons.
    """

    return (
        value.strip()
        .lower()
        .replace(" ", "")
        .replace(".", "")
        .replace("ø", "oe")
        .replace("å", "aa")
        .replace("æ", "ae")
        .replace("?", "")
        .replace("-", "")
        .replace("_", "")
    )


def replace_placeholders(text: str, data: dict) -> str:
    """
    Replace placeholders in the form {key} with values from the data dictionary.
    """

    # ----------------------------------------
    # Fix malformed placeholders
    # ----------------------------------------
    text = re.sub(
        r"\{<[^>]+>(.*?)</[^>]+>\}",
        r"{\1}",
        text
    )

    def repl(match):

        # Extract placeholder key and clean invisible characters
        key = match.group(1).replace("​", "").strip()

        value = data.get(key)

        # If no value exists, keep the original placeholder
        if value is None:
            return match.group(0)

        # Wrap replacements in a blue color span so inserted values are visually distinguishable in the final document.
        return f'<span style="color:#0F9ED5">{value}</span>'

    # Replace all placeholders of the form {key}
    return re.sub(r"\{([^{}]+)\}", repl, text)
